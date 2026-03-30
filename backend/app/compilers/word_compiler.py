"""Word Document Compiler - Compiles agent results into Word files."""
import io
from typing import Dict, Any, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import logging

logger = logging.getLogger(__name__)


class WordCompiler:
    """Compiles agent outputs into Word documents."""
    
    def compile(
        self,
        structure: Dict[str, Any],
        data: List[Dict[str, Any]],
        styling: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Compile all components into a Word document."""
        
        try:
            doc = Document()
            
            # Get structure components
            title = structure.get("title", "Document")
            subtitle = structure.get("subtitle", "")
            sections = structure.get("sections", [])
            tables = structure.get("tables", [])
            
            # Apply document-wide styles
            self._apply_document_styles(doc, styling)
            
            # Add title
            if title:
                title_para = doc.add_heading(title, level=0)
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if subtitle:
                subtitle_para = doc.add_paragraph(subtitle)
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                subtitle_para.runs[0].font.size = Pt(12)
                subtitle_para.runs[0].font.italic = True
            
            # Add sections
            for section in sections:
                self._add_section(doc, section, data, styling)
            
            # Add tables
            for table_config in tables:
                self._add_table(doc, table_config, data)
            
            # Save to bytes
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            return {
                "content": output.getvalue(),
                "filename": "generated_document.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "size": len(output.getvalue())
            }
            
        except Exception as e:
            logger.error(f"Word compilation failed: {e}")
            return self._create_fallback_output()
    
    def _apply_document_styles(self, doc: Document, styling: Dict[str, Any]):
        """Apply document-wide styling."""
        
        paragraph_styles = styling.get("paragraph_styles", {})
        
        # Configure normal style
        normal_style = doc.styles["Normal"]
        normal_font = normal_style.font
        normal_font.name = paragraph_styles.get("Normal", {}).get("font", "Calibri")
        normal_font.size = Pt(paragraph_styles.get("Normal", {}).get("size", 11))
        
        # Configure heading styles
        for i in range(1, 4):
            style_key = f"Heading {i}"
            if style_key in paragraph_styles:
                style_config = paragraph_styles[style_key]
                try:
                    heading_style = doc.styles[style_key]
                    heading_style.font.name = style_config.get("font", "Calibri")
                    heading_style.font.size = Pt(style_config.get("size", 16 - (i-1)*2))
                    heading_style.font.bold = style_config.get("bold", True)
                    
                    color = style_config.get("color", "#000000")
                    if color.startswith("#"):
                        r = int(color[1:3], 16)
                        g = int(color[3:5], 16)
                        b = int(color[5:7], 16)
                        heading_style.font.color.rgb = RGBColor(r, g, b)
                except Exception as e:
                    logger.warning(f"Failed to apply heading style {i}: {e}")
    
    def _add_section(
        self,
        doc: Document,
        section: Dict[str, Any],
        data: List[Dict[str, Any]],
        styling: Dict[str, Any]
    ):
        """Add a section to the document."""
        
        heading = section.get("heading", "")
        level = section.get("level", 1)
        content = section.get("content", "")
        bullet_points = section.get("bullet_points", [])
        
        # Add heading
        if heading:
            doc.add_heading(heading, level=min(level, 3))
        
        # Add content paragraph
        if content:
            para = doc.add_paragraph(content)
            para.paragraph_format.line_spacing = 1.15
        
        # Add bullet points
        for point in bullet_points:
            para = doc.add_paragraph(point, style="List Bullet")
            para.paragraph_format.line_spacing = 1.15
    
    def _add_table(
        self,
        doc: Document,
        table_config: Dict[str, Any],
        data: List[Dict[str, Any]]
    ):
        """Add a table to the document."""
        
        headers = table_config.get("headers", [])
        rows = table_config.get("rows", [])
        
        if not headers and data:
            headers = list(data[0].keys())
            rows = [[str(row.get(h, "")) for h in headers] for row in data[:10]]
        
        if not headers:
            return
        
        # Create table
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = str(header)
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add data rows
        for row_idx, row_data in enumerate(rows):
            row_cells = table.rows[row_idx + 1].cells
            for col_idx, value in enumerate(row_data):
                if col_idx < len(row_cells):
                    row_cells[col_idx].text = str(value)
    
    def _create_fallback_output(self) -> Dict[str, Any]:
        """Create a minimal fallback Word document."""
        doc = Document()
        doc.add_heading("Document Generation Error", 0)
        doc.add_paragraph("There was an error generating your document.")
        doc.add_paragraph("Please try again with a different request.")
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return {
            "content": output.getvalue(),
            "filename": "error.docx",
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "size": len(output.getvalue())
        }


# Global instance
word_compiler = WordCompiler()
